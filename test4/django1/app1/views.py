from django.shortcuts import render
from app1.forms import inputform
from docx import Document
from django.http import HttpResponse
from datetime import datetime
import io
import math
def home(request):
    if request.method=="POST":
        form1=inputform(request.POST)
        if form1.is_valid():
            data=form1.cleaned_data
            dt=datetime.now()
            name=data.get("input0")
            height=data.get("input1")
            weight=data.get("input2")
            IVSd=data.get("input3")
            LVIDd=data.get("input4")
            PWTd=data.get("input5")
            Bio_Gen=data.get("input6")
            #we need to find something LVM(LEFT VENTIRCULAR MASS)
            BSA= round(math.sqrt((height*weight)/3600),2)
            total=IVSd+LVIDd+PWTd
            LVM=(0.8*(1.04*((total**3))-(LVIDd**3)))+0.6 
            LVM=round(LVM,2)
            LVMI= (LVM)/(BSA)
            RS="UNDEFINED"
            if(LVMI<=115 and Bio_Gen=='M'):
                RS="NORMAL"
            elif(LVMI<=95 and Bio_Gen=='F'):
                RS="NORMAL"
            else:
                RS="HIGH / ABNORMAL"

            if 'download' in request.POST:
                doc=Document()
                doc.add_heading("PATIENT DATA",level=1)
                doc.add_paragraph(f"""
                                  -------------------------------------------------------------------------------
                                                            ADITYA HOSPITAL
                                         Dr. Ramesh Chakravarthi – Padubidri Clinic
                                  -------------------------------------------------------------------------------
                              Details:
                                  Patient Name: {name}
                                  Gender: {Bio_Gen}
                                  Age : 20
                                  Date: {dt}
                              The values are: 
                              Patient Height: {height}
                              Patient Weight:{weight}
                              IVSd Value:{IVSd}
                              LVIDd:{LVIDd}
                              PWTd:{PWTd}
                              Gender(Biological) : {Bio_Gen}
                              Results:
                                   (i)Left Ventricular Mass(LVM): ~{LVM}
                                   (ii)LVM Index : ~{LVMI}
                                   (iii)Body Surface Area(BSA):~{BSA}
                              ___________________________
                             Reference Ranges:
                             Normal(Men) : LVMI< 115g/m^2
                             Normal(Women): LVMI <95g/m^2
                              """)
                file_stream=io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                response=HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                 )
                response['Content-Disposition']='attachment; filename="user_data.docx"'
                return response
            return render(request,'app1/index.html',{
            'PName':name,
            'form':form1,
            'LVM':LVM,
            'LVMI':LVMI,
            'BSA':BSA,
            'GEN':Bio_Gen,
            'RS':RS,
            'dt':dt,
            })
    else:
        form1=inputform()
    return render(request,'app1/index.html',{'form':form1})
